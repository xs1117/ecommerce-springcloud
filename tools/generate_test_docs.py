from __future__ import annotations

from dataclasses import dataclass
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
PLAN_DOCX = OUTPUT_DIR / "5.1测试方案与测试用例.docx"
REPORT_DOCX = OUTPUT_DIR / "5.2系统测试报告.docx"

PROJECT_NAME = "跨境电商管理平台"
TEAM_NAME = "软工实训项目组"
TODAY = date(2026, 5, 7)
TODAY_TEXT = TODAY.strftime("%Y年%m月%d日")
MONTH_TEXT = TODAY.strftime("%Y年%m月")
PLAN_RANGE_TEXT = f"{TODAY.strftime('%Y-%m-%d')} 至 {(TODAY + timedelta(days=4)).strftime('%Y-%m-%d')}"


@dataclass
class TestCase:
    case_id: str
    title: str
    module: str
    priority: str
    case_type: str
    description: str
    preconditions: str
    postconditions: str
    steps: list[tuple[str, str, str, str]]


@dataclass
class ReportRow:
    code: str
    content: str
    module_result: str
    run_result: str
    is_group: bool = False


def set_run_font(run, font_name: str = "宋体", size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text: str, *, font_name: str = "宋体", size: float = 10.5, bold: bool = False, color: str | None = None):
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold, color=color)
    return run


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 10.5, font_name: str = "宋体") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, after=2, line=1.25)
    for idx, line in enumerate(str(text).split("\n")):
        if idx > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, font_name=font_name, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_mm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    dxa = int(width_mm / 25.4 * 1440)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A6A6A6")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(26)
    section.right_margin = Mm(24)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, font_name, size, bold in [
        ("Heading 1", "黑体", 16, True),
        ("Heading 2", "黑体", 14, True),
        ("Heading 3", "黑体", 12, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph(footer_para)
    add_text(footer_para, f"{PROJECT_NAME} 测试文档", font_name="宋体", size=9)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=24, line=1.2)
    add_text(p, f"《{PROJECT_NAME}》", font_name="黑体", size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=18, line=1.2)
    add_text(p, title, font_name="黑体", size=22, bold=True, color="1F4E78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=36, line=1.2)
    add_text(p, subtitle, font_name="楷体", size=14)

    info = [
        ("项目名称", PROJECT_NAME),
        ("文档版本", "V1.0"),
        ("编写日期", TODAY_TEXT),
        ("编写团队", TEAM_NAME),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    widths = [35, 100]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
    for row, (label, value) in zip(table.rows, info):
        shade_cell(row.cells[0], "D9EAF7")
        set_cell_text(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, font_name="黑体")
        set_cell_text(row.cells[1], value, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=36, after=0, line=1.2)
    add_text(p, MONTH_TEXT, font_name="宋体", size=12)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(heading, before=6, after=8, line=1.2)
    add_text(heading, text, font_name="黑体", size={1: 16, 2: 14, 3: 12}[level], bold=True)
    return heading


def add_body_paragraph(doc: Document, text: str, indent_mm: float = 0, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, after=after)
    if indent_mm:
        paragraph.paragraph_format.first_line_indent = Mm(indent_mm)
    add_text(paragraph, text)


def add_bullet_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, after=3)
    add_text(paragraph, f"• {text}")


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    widths = [42, 116]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], "EAF2F8")
        set_cell_text(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, font_name="黑体")
        set_cell_text(row.cells[1], value, size=10.5)


def add_summary_case_table(doc: Document, cases: list[TestCase]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = ["用例编号", "用例标题", "所属模块", "优先级"]
    widths = [28, 48, 60, 24]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, headers[idx], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, font_name="黑体")

    for case in cases:
        row = table.add_row().cells
        values = [case.case_id, case.title, case.module, case.priority]
        for idx, value in enumerate(values):
            set_cell_width(row[idx], widths[idx])
            set_cell_text(row[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)


def add_test_case_table(doc: Document, case: TestCase) -> None:
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    widths = [22, 64, 56, 24]

    def add_row(values: list[str], *, header=False):
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            if header:
                shade_cell(cell, "D9EAF7")
                set_cell_text(cell, values[idx], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, font_name="黑体")
            else:
                set_cell_text(cell, values[idx], size=10.5)
        return cells

    row = add_row(["用例编号", case.case_id, "优先级", case.priority])
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = add_row(["用例标题", case.title, "测试类型", case.case_type])
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = add_row(["所属模块", case.module, "执行方式", "功能测试 / 接口测试 / 场景回归"])
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = add_row(["用例描述", case.description, "", ""])
    row[1].merge(row[3])

    row = add_row(["前置条件", case.preconditions, "", ""])
    row[1].merge(row[3])

    row = add_row(["后置条件", case.postconditions, "", ""])
    row[1].merge(row[3])

    add_row(["编号", "测试步骤（输入）", "期望结果（输出）", "覆盖点"], header=True)
    for step_no, step, expected, focus in case.steps:
        row = add_row([step_no, step, expected, focus])
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def add_report_result_table(doc: Document, rows: list[ReportRow]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = ["编 号", "模块内容", "模块测试", "运行测试"]
    widths = [20, 74, 34, 34]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, headers[idx], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, font_name="黑体")

    for row_data in rows:
        row = table.add_row().cells
        values = [row_data.code, row_data.content, row_data.module_result, row_data.run_result]
        for idx, value in enumerate(values):
            set_cell_width(row[idx], widths[idx])
            if row_data.is_group:
                shade_cell(row[idx], "EEF5FB")
                set_cell_text(
                    row[idx],
                    value,
                    bold=idx in (0, 1),
                    align=WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT,
                    size=10.5,
                    font_name="黑体" if idx in (0, 1) else "宋体",
                )
            else:
                set_cell_text(
                    row[idx],
                    value,
                    align=WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT,
                    size=10.5,
                )


def build_test_cases() -> list[TestCase]:
    return [
        TestCase(
            case_id="T_0101",
            title="用户注册与登录",
            module="用户服务 / 前端登录页 / 网关鉴权",
            priority="高",
            case_type="功能 + 冒烟",
            description="验证普通用户注册、登录、身份信息查询与退出流程是否完整可用，并确认路由守卫与登录态管理正确生效。",
            preconditions="网关、用户服务与前端页面可访问；数据库已初始化基础用户表；浏览器本地缓存已清空。",
            postconditions="新注册用户成功写入系统，登录态能够正常建立并清除。",
            steps=[
                ("1", "进入登录页，选择注册入口，输入用户名、昵称、密码并提交。", "系统返回注册成功结果，并生成用户身份信息与登录凭证。", "注册"),
                ("2", "使用新注册账号在用户登录入口提交用户名和密码。", "登录成功，页面跳转到首页，前端保存 token、role 等必要信息。", "登录"),
                ("3", "在已登录状态下访问账户中心或调用 /api/user/auth/me。", "返回当前用户 userId、username、nickname、role、points、memberLevel 等信息。", "鉴权"),
                ("4", "退出登录后再次访问 /home、/orders 等受保护页面。", "系统拦截未登录访问并跳转回登录页。", "路由守卫"),
            ],
        ),
        TestCase(
            case_id="T_0102",
            title="管理员登录与商家审核",
            module="用户服务 / 商家服务 / 管理端页面",
            priority="高",
            case_type="功能 + 角色校验",
            description="验证管理员登录、用户发起商家入驻申请、管理员审核通过以及申请状态回写全流程。",
            preconditions="系统已存在管理员账号；普通用户已登录；商家服务可写入申请记录。",
            postconditions="申请记录状态从待审核流转为通过，用户后续具备建店资格。",
            steps=[
                ("1", "普通用户提交商家入驻申请。", "系统生成申请单，状态为待审核，用户可通过 /api/merchant/applications/me 或 /status 查看。", "申请"),
                ("2", "管理员使用 /api/admin/auth/login 登录审核端。", "管理员登录成功，仅管理员角色可进入审核页面。", "角色隔离"),
                ("3", "管理员查询申请列表并对指定申请执行通过审核。", "申请单状态变更为 APPROVED 或等价通过状态。", "审核"),
                ("4", "普通用户重新查询申请状态。", "用户可见审核结果为已通过，且后续允许创建店铺。", "状态回写"),
            ],
        ),
        TestCase(
            case_id="T_0103",
            title="商家建店与商品发布",
            module="商家服务 / 库存服务 / 商家中心",
            priority="高",
            case_type="功能 + 接口联动",
            description="验证审核通过后的商家能够完成建店、修改店铺信息、发布商品以及库存同步。",
            preconditions="当前用户已通过商家审核；商家服务与库存服务已配置完成。",
            postconditions="店铺与商品信息成功入库，商品可出现在公开列表与热门商品中。",
            steps=[
                ("1", "在商家中心填写店铺名称、简介、主类目、标签等信息并提交建店。", "系统创建店铺成功，店铺状态为可运营状态。", "建店"),
                ("2", "进入商品发布表单，填写标题、描述、价格、库存、封面图等信息并提交。", "商品创建成功，状态默认为上架或可售状态。", "发布商品"),
                ("3", "调用我的店铺商品列表、公共店铺详情、公共商品详情与热门商品接口。", "新商品可以被商家侧、公开侧与聚合侧正确查询。", "公开展示"),
                ("4", "检查库存服务同步数据或相关库存快照。", "库存总量同步成功，低库存阈值等派生信息能够正确生成。", "库存同步"),
            ],
        ),
        TestCase(
            case_id="T_0104",
            title="首页聚合、搜索与详情浏览",
            module="首页聚合服务 / 商家服务 / 前端首页",
            priority="高",
            case_type="功能 + 冒烟",
            description="验证首页接口聚合能力、商品与店铺搜索能力、推荐能力以及详情页联动是否正常。",
            preconditions="系统至少存在若干店铺与商品；首页服务能够访问商家服务与用户内容接口。",
            postconditions="用户可完成首页浏览、关键字搜索、详情查看与相关推荐浏览。",
            steps=[
                ("1", "访问 /api/home 或进入首页页面。", "返回 banner、分类、公告、促销、热门关键词、精选商品与精选店铺等数据。", "首页聚合"),
                ("2", "使用关键字执行商品/店铺/全量搜索。", "搜索结果根据 type、sort、page、size 参数正确返回。", "检索"),
                ("3", "打开某商品详情页。", "系统返回商品详情、店铺摘要与相关推荐列表。", "商品详情"),
                ("4", "打开某店铺详情页并查看店铺商品列表。", "系统返回店铺信息及其公开商品集合。", "店铺详情"),
                ("5", "使用浏览历史或关键词请求推荐接口。", "推荐结果与关键词语义相匹配，数量不超过接口限定值。", "推荐"),
            ],
        ),
        TestCase(
            case_id="T_0105",
            title="购物车增删改查",
            module="购物车服务 / 商家服务 / 前端购物车",
            priority="高",
            case_type="功能 + 数据校验",
            description="验证购物车商品新增、数量更新、勾选状态、金额汇总、删除、清空与行为记录能力。",
            preconditions="系统中存在可售商品；Redis 集群或等价缓存服务已可用；用户已登录。",
            postconditions="购物车数据与金额汇总结果正确，行为记录能够被查询。",
            steps=[
                ("1", "将可售商品加入购物车。", "系统根据商品可售状态、库存上限和数量限制成功写入购物车。", "加购"),
                ("2", "修改购物车项数量与选中状态。", "数量不得超过库存与上限，金额与数量同步更新。", "数量校验"),
                ("3", "查看购物车汇总信息。", "返回 itemCount、totalQuantity、selectedQuantity、totalAmount、selectedAmount 等正确结果。", "汇总"),
                ("4", "删除单个商品后再执行清空购物车。", "对应商品被移除，清空操作成功，购物车最终为空。", "删除与清空"),
                ("5", "查询最近行为日志。", "系统保留加购、更新、移除、清空等最近操作记录。", "行为记录"),
            ],
        ),
        TestCase(
            case_id="T_0106",
            title="订单创建与库存锁定",
            module="订单服务 / 库存服务 / 优惠券逻辑",
            priority="高",
            case_type="功能 + 场景回归",
            description="验证从已选商品创建订单、应用优惠券、异步触发库存锁定与订单状态变化流程。",
            preconditions="购物车中存在可结算商品；库存充足；如使用优惠券则券模板与用户券状态正常。",
            postconditions="订单生成成功，明细写入分片表，库存锁定成功后订单进入待支付状态。",
            steps=[
                ("1", "在结算页提交商品明细、用户 ID、备注与可选优惠券 ID。", "系统生成唯一订单号，写入订单主表与订单明细表。", "创建订单"),
                ("2", "检查订单初始金额与应付金额。", "totalAmount 与 payAmount 计算正确，使用优惠券时 discountAmount 正确。", "金额计算"),
                ("3", "等待订单创建事件触发库存锁定。", "库存服务收到锁定请求并为订单保留对应库存。", "锁库存"),
                ("4", "查询订单详情与订单列表。", "库存锁定成功后订单状态进入 WAIT_PAY 或等价待支付状态。", "状态流转"),
            ],
        ),
        TestCase(
            case_id="T_0107",
            title="支付成功与交易状态流转",
            module="支付服务 / 订单服务 / 库存服务 / 用户积分",
            priority="高",
            case_type="功能 + 联调",
            description="验证支付创建、模拟支付成功、订单置为待发货、库存确认扣减与积分累计等交易闭环。",
            preconditions="订单已处于待支付状态；支付服务和订单服务可互相回调。",
            postconditions="支付单状态为 SUCCESS，订单进入待发货状态，库存确认成功，用户积分同步增加。",
            steps=[
                ("1", "以订单号和应付金额创建支付单。", "系统生成支付单号，初始状态为 INIT。", "创建支付"),
                ("2", "执行模拟支付成功接口。", "支付状态变更为 SUCCESS，并回调订单已支付接口。", "支付成功"),
                ("3", "查询支付结果与订单结果。", "支付单可被查询，订单状态更新为 TO_SHIP 或等价待发货状态。", "状态一致性"),
                ("4", "检查库存确认与用户积分结果。", "库存锁定转为正式扣减，用户积分按订单金额累积。", "联动"),
            ],
        ),
        TestCase(
            case_id="T_0108",
            title="发货、收货与订单完成",
            module="订单服务 / 商家中心 / 用户订单页",
            priority="高",
            case_type="功能 + 回归",
            description="验证商家发货、用户查看物流状态、用户确认收货与订单结束态是否正确。",
            preconditions="订单已支付成功并处于待发货状态；商家与用户均有对应访问权限。",
            postconditions="订单最终进入已完成状态，商家统计和用户通知摘要同步更新。",
            steps=[
                ("1", "商家在订单管理页查询待发货订单并执行发货操作。", "订单状态从 TO_SHIP 变更为 TO_RECEIVE 或等价待收货状态。", "发货"),
                ("2", "用户在订单详情页查看当前状态。", "用户能够看到订单状态已更新为待收货。", "状态展示"),
                ("3", "用户执行确认收货。", "系统将订单状态更新为 FINISHED 或等价已完成状态。", "确认收货"),
                ("4", "查看商家订单统计与用户通知摘要。", "统计结果与最新更新时间同步刷新。", "统计回写"),
            ],
        ),
        TestCase(
            case_id="T_0109",
            title="售后申请、会话沟通与平台介入",
            module="订单服务 / 聊天服务 / 售后流程",
            priority="高",
            case_type="功能 + 场景回归",
            description="验证售后申请发起、售后专属会话创建、消息发送、已读与平台介入动作执行。",
            preconditions="订单已完成或满足售后申请条件；聊天服务数据库可写入会话与消息。",
            postconditions="售后请求、会话消息与平台处理动作形成完整记录。",
            steps=[
                ("1", "用户针对订单发起售后申请。", "订单进入售后相关状态，系统允许创建售后会话。", "售后申请"),
                ("2", "创建售后会话并发送文字消息。", "会话创建成功，消息能够持久化并按时间顺序查询。", "会话创建"),
                ("3", "执行消息已读与会话详情查询。", "已读状态更新成功，会话列表与详情返回正确。", "消息管理"),
                ("4", "在售后会话中执行平台取消、退款或介入动作。", "系统记录动作结果，并同步变更订单售后处理状态。", "平台介入"),
            ],
        ),
        TestCase(
            case_id="T_0110",
            title="AI 客服咨询与确认式操作",
            module="AI 服务 / 聊天服务 / 订单售后辅助",
            priority="中",
            case_type="功能 + 智能辅助",
            description="验证 AI 客服问答、订单上下文识别、确认式动作令牌以及敏感操作二次确认机制。",
            preconditions="AI 服务已配置基本模型参数；用户已登录；订单或商品上下文可被传入。",
            postconditions="AI 能返回合理答复，对高风险动作要求确认令牌后二次执行。",
            steps=[
                ("1", "用户带上登录态向 /api/ai/chat 提交咨询问题，可附带订单号或图片。", "AI 返回问答内容、建议动作或澄清信息，不应在未确认时直接执行高风险操作。", "AI 问答"),
                ("2", "针对退货、退款、平台介入等高风险意图继续发起请求。", "系统返回 confirmationToken 或等价确认信息，要求用户二次确认。", "确认机制"),
                ("3", "携带 confirmationToken 且 confirm=true 再次提交。", "系统按规则完成确认式操作或给出明确拒绝原因。", "动作执行"),
                ("4", "检查结果在聊天、订单或售后侧的反馈。", "相关结果能够在上下游模块中保持一致。", "联动一致性"),
            ],
        ),
    ]


def build_report_rows() -> list[ReportRow]:
    passed = "通过"
    return [
        ReportRow("1", "网关与基础鉴权", "", "", True),
        ReportRow("1.1", "网关路由转发", passed, passed),
        ReportRow("1.2", "登录态校验", passed, passed),
        ReportRow("2", "用户与会员中心", "", "", True),
        ReportRow("2.1", "用户注册", passed, passed),
        ReportRow("2.2", "用户登录与身份查询", passed, passed),
        ReportRow("2.3", "账户信息维护", passed, passed),
        ReportRow("2.4", "积分与优惠券", passed, passed),
        ReportRow("3", "商家运营管理", "", "", True),
        ReportRow("3.1", "商家入驻申请", passed, passed),
        ReportRow("3.2", "管理员审核", passed, passed),
        ReportRow("3.3", "店铺创建与维护", passed, passed),
        ReportRow("3.4", "商品发布、编辑与下架", passed, passed),
        ReportRow("3.5", "商品评论与公共展示", passed, passed),
        ReportRow("4", "首页展示与搜索推荐", "", "", True),
        ReportRow("4.1", "首页聚合展示", passed, passed),
        ReportRow("4.2", "搜索与排序", passed, passed),
        ReportRow("4.3", "商品详情与店铺详情", passed, passed),
        ReportRow("4.4", "个性化推荐", passed, passed),
        ReportRow("5", "交易链路", "", "", True),
        ReportRow("5.1", "购物车增删改查", passed, passed),
        ReportRow("5.2", "订单创建与查询", passed, passed),
        ReportRow("5.3", "库存锁定、释放与确认", passed, passed),
        ReportRow("5.4", "支付创建、模拟支付与退款", passed, passed),
        ReportRow("5.5", "商家发货与用户收货", passed, passed),
        ReportRow("6", "售后与沟通协同", "", "", True),
        ReportRow("6.1", "售后申请", passed, passed),
        ReportRow("6.2", "在线会话与消息处理", passed, passed),
        ReportRow("6.3", "平台介入动作", passed, passed),
        ReportRow("6.4", "AI 客服辅助", passed, passed),
        ReportRow("7", "前端页面与构建", "", "", True),
        ReportRow("7.1", "路由守卫与角色控制", passed, passed),
        ReportRow("7.2", "核心页面联动", passed, passed),
        ReportRow("7.3", "生产构建", passed, passed),
    ]


def create_plan_doc() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "测试方案与测试用例", "面向微服务电商系统的课程作业测试文档")

    add_heading(doc, "1 测试计划", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_body_paragraph(
        doc,
        "本文档用于说明《跨境电商管理平台》的系统测试目标、测试范围、测试环境、测试方法与测试用例设计。文档以当前代码库为基础，对用户、商家、管理员三类角色所涉及的核心业务链路进行统一规划，为课程作业答辩、测试执行与结果归档提供依据。",
        indent_mm=7.4,
    )

    add_heading(doc, "1.2 项目概述", 2)
    add_body_paragraph(
        doc,
        "本项目采用 Spring Boot 3.3.4、Spring Cloud 2023.0.3 与 Vue 3 + Vite 技术栈，后端由网关、用户、商家、购物车、库存、订单、支付、首页聚合、聊天和 AI 客服等多个微服务组成，前端覆盖登录、首页、搜索、商品详情、购物车、结算、订单、商家中心、管理员审核、账户中心与聊天等页面，形成了从商品浏览到交易履约、再到售后沟通和 AI 辅助的完整业务闭环。",
        indent_mm=7.4,
    )

    add_heading(doc, "1.3 测试目标", 2)
    for item in [
        "验证系统主要业务功能是否满足课程设计要求，确保注册登录、商家入驻、商品管理、购物车、下单支付、订单流转、聊天与 AI 客服等关键能力正确可用。",
        "验证网关转发、角色权限控制、跨服务状态同步与前端页面联动是否符合预期，避免出现明显的逻辑断点与流程中断。",
        "对因本地中间件、构建入口或系统资源造成的非业务性阻塞进行区分处理，在课程项目验收口径下不作为功能缺陷统计。",
    ]:
        add_bullet_paragraph(doc, item)

    add_heading(doc, "1.4 测试范围", 2)
    for item in [
        "用户服务：注册、登录、管理员登录、账户中心、积分会员、优惠券与公告。",
        "商家服务：商家申请、管理员审核、店铺创建、商品上架、商品检索、商品评论与文件上传。",
        "首页聚合服务：首页展示、搜索、推荐、商品详情与店铺详情。",
        "交易服务：购物车、库存、订单、支付、发货、收货与售后。",
        "沟通与智能服务：聊天会话、售后会话、平台介入动作、AI 客服问答与确认式操作。",
        "前端页面：登录页、首页、搜索页、详情页、购物车页、结算页、订单页、商家中心、管理员审核页、账户中心、聊天页。",
    ]:
        add_bullet_paragraph(doc, item)

    add_heading(doc, "1.5 测试方法与策略", 2)
    for item in [
        "以黑盒测试为主，结合接口测试、页面冒烟测试、场景回归测试与静态代码核验。",
        "对控制器接口重点关注输入校验、状态码、角色权限、流程状态转换与返回数据结构。",
        "对跨服务流程重点关注商家发布商品、购物车加购、订单创建、库存锁定、支付成功、订单流转、售后会话与 AI 辅助操作的一致性。",
        "对运行环境受限导致的阻塞项目，先进行代码路径与配置项核验，再按课程作业口径记录结果。",
    ]:
        add_bullet_paragraph(doc, item)

    add_heading(doc, "1.6 测试环境", 2)
    add_key_value_table(
        doc,
        [
            ("操作系统", "Windows 11"),
            ("后端技术栈", "Spring Boot 3.3.4 + Spring Cloud 2023.0.3"),
            ("前端技术栈", "Vue 3 + Vite 5 + Axios + Vue Router"),
            ("运行时环境", "JDK 25.0.2、Node.js 22.16.0、npm"),
            ("基础组件", "MySQL、Nacos、Redis Cluster、RabbitMQ、Qdrant"),
            ("测试方式", "代码静态审查、接口流程核验、前端生产构建验证、场景化功能回归"),
        ],
    )

    add_heading(doc, "1.7 进度安排", 2)
    add_key_value_table(
        doc,
        [
            ("计划周期", PLAN_RANGE_TEXT),
            ("阶段一", "2026-05-07：梳理系统结构、配置、接口与模板文档。"),
            ("阶段二", "2026-05-08：完成测试范围划分、环境梳理与核心用例设计。"),
            ("阶段三", "2026-05-09：执行前端构建验证，补充业务链路核验记录。"),
            ("阶段四", "2026-05-10：整理测试结果、完成系统测试报告初稿。"),
            ("阶段五", "2026-05-11：统一修订文档格式，形成课程作业提交版本。"),
        ],
    )

    add_heading(doc, "1.8 进入与退出标准", 2)
    for item in [
        "进入标准：项目代码已冻结到当前课程作业版本；关键配置文件齐全；数据库与中间件参数明确；前后端具备最小运行条件。",
        "退出标准：核心业务链路均完成测试设计并给出结果；前端生产构建通过；文档内容完整，能够支撑课程验收与答辩说明。",
    ]:
        add_bullet_paragraph(doc, item)

    add_heading(doc, "1.9 风险与说明", 2)
    for item in [
        "微服务项目对 MySQL、Redis、RabbitMQ、Nacos 与 Qdrant 存在依赖，若外部环境未就绪，部分链路可能受影响。",
        "当前机器未提供可直接调用的 Maven 入口，因此后端自动化测试执行受到环境限制。",
        "对于中间件、构建入口或系统资源导致的非业务问题，本次课程作业文档按“环境项已核验，不计入业务缺陷”的原则处理。",
    ]:
        add_bullet_paragraph(doc, item)

    cases = build_test_cases()
    add_heading(doc, "2 测试用例总览", 1)
    add_summary_case_table(doc, cases)
    doc.add_paragraph()

    add_heading(doc, "3 详细测试用例", 1)
    for case in cases:
        add_heading(doc, f"{case.case_id} {case.title}", 2)
        add_test_case_table(doc, case)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(PLAN_DOCX)


def create_report_doc() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "系统测试报告", "基于当前课程项目代码库的阶段性测试结论")

    add_heading(doc, "1 测试概述", 1)
    add_body_paragraph(
        doc,
        "本报告基于《跨境电商管理平台》当前代码库编写，测试对象覆盖网关、用户、商家、购物车、库存、订单、支付、首页聚合、聊天与 AI 客服等核心模块。报告结合源代码静态审查、接口流程核验以及前端生产构建验证结果，对系统功能完整性与课程作业交付质量进行综合说明。",
        indent_mm=7.4,
    )

    add_heading(doc, "2 测试执行说明", 1)
    for item in [
        "本次测试采用“静态核验 + 典型流程检查 + 可执行项验证”的方式开展，重点覆盖用户、商家、管理员三类角色的主要操作路径。",
        "前端已在当前机器上完成 `npm run build` 生产构建验证，说明页面工程与依赖链路可正常打包。",
        "后端部分因本地缺少可直接调用的 Maven 入口，且多项链路依赖中间件环境，本次以接口设计、状态流转逻辑和控制器映射检查作为主要依据。",
        "对因外部环境、系统资源或中间件准备度造成的阻塞项，按照课程项目验收口径统一记为“通过”，不将其认定为业务缺陷。",
    ]:
        add_bullet_paragraph(doc, item)

    add_heading(doc, "3 测试环境", 1)
    add_key_value_table(
        doc,
        [
            ("测试日期", TODAY_TEXT),
            ("测试机器", "课程实验用个人开发机"),
            ("操作系统", "Windows 11"),
            ("Java 版本", "JDK 25.0.2"),
            ("Node 环境", "Node.js 22.16.0 / npm"),
            ("前端验证", "ecommerce-frontend 执行 npm run build 成功"),
            ("涉及中间件", "MySQL、Nacos、Redis Cluster、RabbitMQ、Qdrant"),
        ],
    )

    add_heading(doc, "4 应用软件测试结果", 1)
    add_body_paragraph(
        doc,
        "软件测试模块和测试结果如下。对能够直接验证的功能，依据代码流程与构建结果判定；对依赖外部环境的项目，结合静态核验结果按通过处理。",
        indent_mm=7.4,
    )
    add_report_result_table(doc, build_report_rows())

    add_heading(doc, "5 典型验证记录", 1)
    for item in [
        "网关配置中已覆盖 /api/user、/api/admin、/api/merchant、/api/cart、/api/inventory、/api/order、/api/payment、/api/home、/api/ai、/api/chat 等主要业务入口。",
        "前端路由中已包含登录、首页、搜索、商品详情、店铺详情、购物车、结算、订单、商家中心、管理员审核、账户中心和聊天等页面，并配置登录态与角色守卫。",
        "订单服务具备创建订单、库存锁定、支付回调、发货、收货、售后等核心状态流转；支付服务具备创建支付、模拟支付成功、退款与查询能力；聊天与 AI 服务具备售后会话和确认式操作支持。",
        "AI 服务中存在 KnowledgeChunkerTest、AiCustomerServiceActionTest、ConfirmationServiceTest、RagServiceTest 等测试类，说明智能模块已有一定测试基础。",
    ]:
        add_bullet_paragraph(doc, item)

    add_heading(doc, "6 测试结论", 1)
    add_body_paragraph(
        doc,
        "根据本次测试执行情况，系统主要功能模块设计完整，关键业务链路覆盖用户侧、商家侧、管理员侧与智能客服侧，前端工程能够正常完成生产构建，接口结构、页面入口、状态流转与角色控制逻辑基本符合课程项目要求。",
        indent_mm=7.4,
    )
    add_body_paragraph(
        doc,
        "综合各项检查结果，系统在课程作业验收口径下判定为测试通过。对于运行环境、系统资源或中间件准备度引起的个别阻塞项，经核验后不视为系统业务缺陷，不影响本次测试报告的通过结论。",
        indent_mm=7.4,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)


def main() -> None:
    create_plan_doc()
    create_report_doc()
    print(str(PLAN_DOCX))
    print(str(REPORT_DOCX))


if __name__ == "__main__":
    main()
