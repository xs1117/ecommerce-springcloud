from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


WORK_DIR = Path(r"C:\Users\86152\report_work")
TEMPLATE_PATH = WORK_DIR / "template.docx"
OUTPUT_PATH = WORK_DIR / "course_report_final.docx"

PROJECT_NAME = "跨境电商管理平台"
TEAM_NAME = "软工实训项目组"
DOC_CODE = "dhu-siis2026-软工实训项目组"
REPORT_DATE = "2026-05-17"

TEAM_ROWS = [
    ["", "薛硕", "规划总结项目，后端开发，前端开发；负责订单服务、库存服务、AI服务", "25"],
    ["", "张辰远", "项目测试，后端开发；负责购物车服务、主页服务、聊天服务", "25"],
    ["", "陈星宇", "需求规格设计，后端开发，前端开发；负责用户服务、网关服务、前端页面", "25"],
    ["", "夏雯", "概要设计，详细设计，后端开发，前端开发；负责商家服务、支付服务、前端页面", "25"],
]

REVISION_ROWS = [
    ["1", "初稿编写", "V1.0", TEAM_NAME, REPORT_DATE, "完成课程大报告正文与表格内容整理"],
    ["2", "内容复查", "V1.1", TEAM_NAME, REPORT_DATE, "统一模板口径、修正章节内容并完成最终复核"],
]

FUNCTION_TABLE_ROWS = [
    ["用户与账户模块", "注册登录、账户中心、积分与优惠券", "支撑普通用户身份认证、资料维护、积分会员和优惠券兑换使用", "高"],
    ["商家与交易模块", "入驻审核、店铺商品、购物车、订单、支付、库存、发货收货", "完成从商家入驻到用户下单、支付、履约和售后的核心交易链路", "高"],
    ["沟通与智能模块", "聊天会话、AI客服、图搜商品、确认式售后", "支撑用户与商家沟通、AI知识问答、图片搜商品和售后辅助处理", "高"],
]

USE_CASE_TABLE_ROWS = {
    "用例名称": "用户购物与售后闭环",
    "功能简述": "用户从登录、浏览商品、加入购物车、下单支付到售后会话与 AI 辅助处理的完整业务流程",
    "用例编号": "UC-01",
    "执行者": "普通用户、商家、管理员、AI客服",
    "前置条件": "用户已登录，商家已上架商品，网关与基础中间件配置完成",
    "后置条件": "订单完成履约或进入售后流程，会话与处理记录持久化",
    "涉众利益": "平台验证核心业务链路完整性，用户获得购物与售后服务，商家完成订单处理",
    "基本路径": "1. 登录进入首页；2. 搜索并查看商品；3. 加入购物车并提交订单；4. 完成支付；5. 商家发货；6. 用户收货；7. 发起售后并进入会话；8. AI 辅助回答或触发确认式动作",
    "扩展路径": "1. 库存不足时限制下单；2. 支付失败时保留待支付订单；3. AI 无法直接执行时转人工或平台介入",
    "字段列表": "用户ID、商品ID、店铺ID、订单号、支付单号、会话ID、售后状态、confirmationToken",
    "设计规则": "Gateway 统一入口，JWT 鉴权，订单/库存通过事件与锁保证一致性，高风险售后动作必须二次确认",
    "未解决的问题": "可进一步补充更细粒度监控、自动化回归和大规模性能压测",
    "备注": "本用例用于课程演示与系统验收的主链路说明",
}

SECTION_CONTENT = {
    "项目背景": [
        "跨境电商业务同时面向普通用户、商家与平台管理员，既要覆盖商品浏览、交易履约、售后沟通等基础能力，又要处理多角色协同、跨服务调用与数据一致性问题。传统单体电商系统在扩展新功能、承接高并发访问和接入智能客服能力时，往往会出现模块耦合高、维护成本大、业务演进慢等问题，因此需要采用更适合拆分演进的微服务架构方案。",
    ],
    "项目定位": [
        "本项目定位为面向课程实训的综合型跨境电商管理平台，以“用户购物、商家运营、平台审核、售后协同、AI 辅助服务”五类核心场景为主线，构建一个具备完整业务闭环、可演示、可扩展、可持续迭代的微服务项目。",
    ],
    "应用场景": [
        "系统主要应用于用户商品搜索与购买、商家店铺与商品运营、管理员审核与平台介入、订单履约与售后处理、AI 客服知识问答与图片搜商品等场景，能够支撑从前台消费到后台运营再到智能客服协同的完整流程。",
    ],
    "目标人群": [
        "目标人群包括普通购物用户、入驻商家、平台管理员以及课程答辩与验收教师。其中普通用户关注购物体验与售后效率，商家关注店铺经营与订单处理，管理员关注平台治理与审核能力，教师则重点考察系统架构设计、业务完整性与技术实现水平。",
    ],
    "项目方案": [
        "项目采用 Vue 3 + Vite 构建前端界面，后端以 Spring Boot、Spring Cloud 为基础拆分网关、用户、商家、购物车、库存、订单、支付、首页聚合、聊天与 AI 服务十个核心模块，结合 MySQL、Redis Cluster、RabbitMQ、Nacos 与 Qdrant 完成数据存储、缓存、消息解耦、配置治理与向量检索能力建设。",
    ],
    "项目目标": [
        "项目目标是完成一个具备完整电商主链路与智能客服亮点的课程实训系统：一方面实现注册登录、店铺运营、下单支付、履约售后等核心功能，另一方面通过 RAG、多 Agent、图片搜商品与确认式动作机制体现系统在智能化方向上的综合能力。",
    ],
    "项目价值": [
        "项目价值体现在三个方面：其一，形成了一个可提交、可答辩、可联调的综合课程成果；其二，通过清晰的微服务边界与基础设施接入锻炼了工程化开发能力；其三，将 AI 服务嵌入真实电商业务流程，提升了项目的创新性、展示性与后续扩展空间。",
    ],
    "最终呈现形式": [
        "最终成果以源代码仓库、前端可构建工程、后端微服务模块、Docker 中间件配置、测试文档、综合报告和答辩材料的形式呈现。其中前端工程可直接执行构建，后端各服务具备独立模块结构与配置文件，整体满足课程项目提交与展示要求。",
    ],
    "主要功能描述": [
        "系统功能覆盖用户注册登录、账户中心、积分与优惠券、商家入驻与店铺管理、商品发布与评论、首页聚合与搜索推荐、购物车、下单支付、库存锁定、商家发货、用户收货、售后申请、聊天会话、平台介入、AI 问答、图搜商品与确认式售后动作执行等主要模块。",
    ],
    "运行环境": [
        "系统运行环境为 Windows 11，本地使用 Java 25.0.2、Node.js 22.16.0 与 npm 11.4.2 进行开发与构建；后端依赖 MySQL、Nacos、Redis Cluster、RabbitMQ 与 Qdrant，中间件均已提供本地联调用配置，前端默认通过 Vite 在 5173 端口运行，网关统一由 8080 端口对外暴露接口。",
    ],
    "验收标准": [
        "验收以业务链路完整性、页面可达性、接口结构一致性和文档规范性为标准。普通用户需能够完成登录、浏览、下单、支付与售后，商家需能够完成入驻、建店、发货与商品管理，管理员需能够完成审核与平台介入，AI 侧需能够完成问答、检索与确认式动作辅助；同时前端构建通过、模板文档完整即视为达到课程验收要求。",
    ],
    "关键问题": [
        "项目开发中的关键问题主要包括多服务之间的注册发现与路由治理、订单支付库存之间的数据一致性、聊天与售后流程的状态同步、AI 输出与真实业务动作之间的安全边界、商品图片检索与知识问答的落地方式，以及多人协作下前后端与文档口径统一的问题。",
    ],
    "进度安排": [
        "项目按“需求分析与分工、架构搭建与模块拆分、前后端联调、交易与售后闭环完善、AI 能力接入、测试与文档收口”六个阶段推进。团队成员围绕用户、商家、交易、AI 与前端页面分工开发，最后统一完成系统测试、课程报告和答辩材料整理。",
    ],
    "开发预算": [
        "本项目作为课程实训系统，主要使用开源框架、已有开发设备和本地容器化环境，不额外产生商业授权费用。预算重点体现在成员时间投入、联调成本、测试整理和文档编写工作量上，整体符合课程项目低成本实现的特点。",
    ],
    "技术可行性分析": [
        "项目所选技术栈成熟稳定，Spring Boot、Spring Cloud、Vue 3、Redis、RabbitMQ、Nacos 等方案均具有完善的社区生态与工程实践基础。当前仓库已经具备清晰的模块划分、可运行的前端工程、完整的配置文件及多项 AI 相关测试类，说明项目在技术实现上具备较高可行性。",
    ],
    "资源可行性分析": [
        "团队现有资源能够满足项目实现需求。代码仓库已提供前后端工程、数据库初始化脚本、Docker 中间件配置、技术分析文档和答辩材料；团队共有四名成员并已明确模块分工，能够覆盖需求、设计、开发、测试和文档整理等关键工作。",
    ],
    "市场可行性分析": [
        "虽然本项目以课程实训为主要目标，但其业务场景与当前电商平台的智能化趋势高度契合。平台将交易主链路与 AI 客服、图搜商品、售后辅助结合，具备良好的展示价值、教学价值和进一步扩展为原型系统的潜力，因此在应用展示和场景验证层面具有较强可行性。",
    ],
    "数据需求": [
        "系统数据需求覆盖用户域、商家域、交易域、聊天域与 AI 域，既包含较稳定的基础资料，也包含在业务执行过程中不断产生和变化的动态数据，需要通过关系型数据库、缓存与向量库组合存储。",
    ],
    "静态数据": [
        "静态数据主要包括用户账号基础信息、积分与优惠券模板、商城公告、商家申请记录、店铺信息、商品基础信息、分类标签、知识库文档以及商品图片索引来源数据。这类数据更新频率相对较低，但对系统展示与业务执行具有基础支撑作用。",
    ],
    "动态数据": [
        "动态数据主要包括购物车商品项、订单与订单明细、支付单、库存预留记录、聊天会话与消息、售后状态、AI 多轮对话上下文、确认令牌与检索结果等。这部分数据在交易和交互过程中持续变化，对系统一致性与响应效率要求较高。",
    ],
    "数据词典": [
        "核心数据表包括用户服务中的 sys_user、sys_points_ledger、coupon_template、user_coupon、mall_notice，商家服务中的 merchant_application、merchant_store、merchant_product、merchant_product_comment，订单服务中的 order_info_0/order_info_1 与 order_item_0/order_item_1，支付服务中的 payment_order，库存服务中的 inventory_sku、inventory_reservation，聊天服务中的 chat_conversation、chat_message。除关系表外，购物车服务使用 Redis Hash 与 List 保存商品项和行为记录，AI 服务使用 Qdrant 保存知识向量与商品图片索引。",
    ],
    "数据采集": [
        "数据来源包括用户注册与页面交互输入、商家主动维护的店铺与商品信息、管理员审核与配置操作、系统在下单支付履约过程中的自动生成记录，以及 AI 服务对知识文档和商品图片进行向量化处理后的检索索引。各类数据通过页面表单、接口请求、消息事件与服务内部处理共同完成采集。",
    ],
    "功能需求": [
        "功能需求围绕“多角色协同 + 完整交易链路 + 智能客服增强”展开，要求系统既能完成传统电商平台的基本业务，又能体现多服务解耦与 AI 场景落地能力。核心功能模块如下表所示。",
    ],
    "时间特性": [
        "系统对页面交互和接口响应的要求以课程演示流畅为基准，登录、首页查询、购物车计算、订单创建、消息发送等核心操作应在可接受的秒级范围内完成；库存锁定、订单状态更新和消息通知等后台处理需保持较快反馈，避免明显阻塞主链路。",
    ],
    "适应性": [
        "系统需适配用户端、商家端和管理员端的不同访问场景，并支持本地开发、联调展示与后续扩展。由于项目采用微服务拆分、外部化配置和容器化中间件部署，能够较好适应模块独立演进、服务扩缩容和 AI 能力持续增强的需要。",
    ],
    "界面需求": [
        "界面层需要覆盖登录页、首页、搜索结果页、商品详情页、店铺详情页、购物车页、结算页、订单列表与详情页、商家申请页、商家中心页、管理员审核页、账户中心页和聊天页。页面布局需保证角色入口清晰、业务流转顺畅，并为 AI 浮层、购物车抽屉等交互组件预留展示空间。",
    ],
    "接口需求": [
        "系统接口需求包括前端对外访问的 HTTP REST 接口以及服务之间的内部调用接口与消息事件接口。接口必须保证路径清晰、参数规范、鉴权统一、状态可追踪，并满足不同角色下的访问控制要求。",
    ],
    "硬件接口": [
        "本系统不依赖专用硬件设备，主要运行在通用 PC 环境和常规服务器资源之上。硬件层面只要求具备基本网络连接、磁盘存储与可运行 Java、Node、MySQL 及中间件容器的计算能力。",
    ],
    "软件接口": [
        "软件接口分为两类：一类是 Gateway 对前端开放的 /api/user、/api/admin、/api/merchant、/api/cart、/api/order、/api/payment、/api/home、/api/chat、/api/ai 等入口；另一类是微服务内部通过 RestClient/HTTP、RabbitMQ 事件、Redis 访问与 Qdrant 检索形成的内部接口，用于完成库存、订单、聊天和 AI 之间的协同调用。",
    ],
    "其他需求": [
        "系统还需满足安全性、可维护性、可扩展性与可演示性要求。安全上要求统一 JWT 鉴权并对高风险动作增加确认机制；维护上要求模块边界清晰、配置集中；扩展上要求支持新增服务与新 Agent；演示上要求页面入口完整、业务链路可顺畅展示。",
    ],
    "处理流程": [
        "系统主流程分为两条主线：一条是用户购物流程，即登录首页、搜索商品、加入购物车、提交订单、完成支付、商家发货、用户收货及必要时发起售后；另一条是商家运营与平台治理流程，即用户申请入驻、管理员审核、商家建店、发布商品并处理订单、售后与会话。AI 服务贯穿在商品咨询、订单问答、图搜商品和售后动作确认等环节中，作为辅助能力嵌入业务主链路。",
    ],
    "总体结构设计": [
        "总体结构采用“前端应用 + 网关层 + 业务微服务层 + 基础设施层”四层架构。前端使用 Vue 3 统一承载多角色页面，Gateway 负责统一入口与跨域处理，业务层按用户、商家、订单、支付、库存、聊天和 AI 等领域划分服务，基础设施层由 MySQL、Redis Cluster、RabbitMQ、Nacos 与 Qdrant 组成，分别承担数据存储、缓存、解耦、配置治理和向量检索功能。",
    ],
    "功能设计": [
        "功能设计按领域拆分：用户服务负责认证、账户与积分优惠券；商家服务负责入驻、店铺、商品与评论；首页聚合服务负责首页展示与搜索；购物车服务负责购物车与行为记录；订单、支付、库存三服务共同承担交易履约；聊天服务负责普通会话与售后会话；AI 服务负责知识问答、售后意图识别、图搜商品和确认式动作执行。",
    ],
    "数据流转设计": [
        "数据流转以 Gateway 为入口。前端请求先进入网关，再按路径转发至目标服务；涉及主链路状态变化的流程以同步调用与异步消息结合的方式完成，例如下单时同步校验与库存锁定，支付成功后通过事件推动订单与库存状态更新；聊天与 AI 会话则通过上下文拼装、知识检索和服务回写保持流程连贯。Redis 主要承担缓存和轻量状态存储，Qdrant 主要承担向量检索。",
    ],
    "用户界面设计": [
        "用户界面采用角色分层设计。普通用户侧强调首页、搜索、商品详情、购物车、订单和聊天的顺畅跳转；商家侧突出申请入驻、建店、商品管理、订单处理与会话处理；管理员侧聚焦商家审核和平台治理。页面通过统一导航、卡片式商品展示和抽屉式辅助交互完成整体界面组织，兼顾展示性与操作效率。",
    ],
    "数据结构设计": [
        "数据结构采用关系数据与非关系数据结合的方案。关系型数据负责用户、商家、订单、支付、库存、聊天会话等核心业务实体；购物车与行为轨迹采用 Redis 的 Hash 和 List 结构保存，适合高频读写；AI 知识库与商品图片检索索引采用 Qdrant 保存向量数据，支持 TopK 检索与相似度召回。",
    ],
    "接口设计": [
        "接口设计遵循 REST 风格和微服务内部协同规则，保证对外接口清晰、对内接口稳定，并通过统一鉴权和状态码规范保证前后端联调效率。接口按角色、领域和同步/异步方式进行划分。",
    ],
    "外部接口": [
        "外部接口由网关统一暴露，包括用户登录注册与账户接口、管理员登录与审核接口、商家申请建店与商品管理接口、购物车接口、订单与支付接口、首页与搜索接口、聊天接口以及 AI 对话接口。前端通过 Axios 统一访问这些接口，并在请求头中携带 Bearer Token 实现身份校验。",
    ],
    "内部接口": [
        "内部接口主要包括服务间 HTTP 调用与 RabbitMQ 消息事件。购物车服务需要查询商品可售状态，订单服务需要调用库存与商家服务，支付服务需要回写订单状态，聊天服务需要读取订单与商家信息，AI 服务则需要接入聊天、订单与商家上下文；其中库存预留、支付成功等跨服务状态变更通过消息事件实现解耦。",
    ],
    "错误/异常处理设计": [
        "系统在控制器层、服务层和 AI 动作层均考虑了错误与异常处理。对参数非法、身份缺失、状态冲突、库存不足、未找到资源和确认令牌失效等情况，系统需给出明确的反馈，并尽量避免异常向用户直接暴露内部实现细节。",
    ],
    "错误/异常输出信息": [
        "对外错误信息以标准 HTTP 状态码和清晰提示语为主，例如登录失败、未登录、权限不足、库存不足、订单状态不允许当前操作、会话不存在、确认令牌失效等。AI 场景下还需要区分“仅给出建议”“需要二次确认”“无法执行需平台介入”等不同结果类型。",
    ],
    "错误/异常处理对策": [
        "异常处理对策包括：前端进行基本表单校验；后端在控制器和服务层做状态检查；订单与库存通过分布式锁、消息和状态枚举降低并发错误；AI 高风险动作通过 confirmationToken 二次确认；异步任务依赖 RabbitMQ 及消费者处理机制降低主链路阻塞与失败扩散风险。",
    ],
    "系统配置策略": [
        "系统配置采用本地 application.yml 与 Nacos 配置中心结合的方式。公共配置放在 ecommerce-common.yaml 中，各服务再通过各自的 yaml 文件补充数据库、Redis、RabbitMQ、Qdrant 与服务发现参数；AI 相关模型、向量库和检索参数支持通过环境变量覆盖，便于不同环境下灵活调整。",
    ],
    "系统部署方案": [
        "系统部署采用本地联调友好的分层部署方案：MySQL、Redis Cluster、RabbitMQ、Qdrant 和 Nacos 作为基础环境先启动，再依次启动用户、商家、首页聚合、订单、支付、库存、聊天、AI 与网关服务，最后运行前端工程进行访问。仓库已提供 Docker Compose 配置，可用于标准化中间件部署与演示环境复现。",
    ],
    "跨端应用架构设计": [
        "本项目以 Web 前端为核心，采用 Vue 3 单页应用统一承载 PC 端与移动浏览器访问场景。通过响应式布局、抽屉组件与统一 API 接口，系统可在不单独开发原生 App 的情况下完成跨端访问与业务展示，因此跨端架构以 H5 适配为主。",
    ],
    "其他相关技术与方案": [
        "除基础微服务架构外，系统还采用 Redis Cluster 支撑高频访问场景，RabbitMQ 支撑异步解耦，Qdrant 支撑知识向量与图片索引检索，Spring AI 支撑模型调用与上下文编排，并在图搜商品场景中结合 dHash 粗筛与视觉识别兜底，构成项目的重要技术亮点。",
    ],
    "数据库设计": [
        "数据库设计按业务域拆分，用户服务使用 ecommerce_user 保存用户、积分、优惠券与公告等数据；商家服务使用 ecommerce_merchant 保存申请、店铺、商品与评论；订单、支付、库存、聊天分别拥有各自独立数据库或表集合，避免跨域直接写库。订单服务采用 order_info_0/order_info_1 与 order_item_0/order_item_1 的分表方案，便于后续扩展与热点隔离。",
        "在非关系型数据方面，购物车服务使用 Redis Hash 保存购物车商品项、使用 Redis List 保存行为记录，能够满足高频读写和快速汇总需求；AI 服务使用 Qdrant 分别保存知识库向量与商品图片索引，用于 RAG 问答和图搜商品；整体数据库设计兼顾了事务型业务处理、缓存性能和向量检索需求。",
    ],
    "手机端侧部署设计(如有)": [
        "本项目未单独开发原生移动端应用，手机端访问以移动浏览器打开前端 H5 页面为主，因此无需额外的安装包发布流程或专用移动端后端接口。模板中的手机端部分按 H5 适配方案说明。",
    ],
    "手机环境需求": [
        "手机端使用现代浏览器即可访问系统前端页面，要求设备具备基本网络连接能力并支持常见 H5、JavaScript 与图片上传功能。由于系统接口统一由 Gateway 暴露，因此手机端与 PC 端共用同一套业务接口和鉴权机制。",
    ],
    "功能描述": [
        "交易履约与购物车模块负责承接用户从商品详情页进入购物车、修改数量、提交订单、完成支付、触发库存扣减与订单状态流转的完整过程，是系统交易主链路的核心模块。",
    ],
    "性能描述": [
        "该模块需要支持高频读写与较快反馈，购物车汇总、加购与数量修改应及时更新，订单创建和支付确认过程应保持稳定；库存锁定与状态回写则通过锁与事件机制保证在并发场景下的正确性。",
    ],
    "输入": [
        "模块输入主要包括用户身份信息、商品与店铺信息、购物车商品项、结算参数、优惠券选择、订单提交请求、支付动作请求以及库存可售状态等数据。",
    ],
    "输出": [
        "模块输出包括购物车汇总结果、订单编号、订单明细、支付单信息、订单状态、库存预留结果、发货与收货结果以及面向前端展示的金额统计数据。",
    ],
    "程序逻辑": [
        "程序逻辑上先由用户在商品详情或列表页发起加购，购物车服务将商品项写入 Redis；结算时由订单服务组装订单明细并请求库存服务锁定库存；支付完成后由支付服务回写订单状态并触发后续履约；商家发货、用户收货与售后申请都围绕订单状态枚举进行合法性控制。",
    ],
    "限制条件": [
        "该模块受商品可售状态、库存数量、订单状态流转规则、优惠券使用条件和支付确认结果约束；当库存不足、订单状态不匹配或支付结果未完成时，系统必须阻止后续非法操作。",
    ],
    "项目总结": [
        "本项目围绕跨境电商核心业务与 AI 增强能力完成了较完整的系统设计与实现，具备清晰的服务边界、较完整的页面链路和较强的课程展示价值，达到了综合实训项目的总体目标。",
    ],
    "经验成果": [
        "通过本次项目实践，团队在需求分析、架构拆分、跨模块协同开发、联调测试和课程文档编写等方面形成了较系统的工程经验。",
    ],
    "成员1收获": [
        "团队对微服务架构下的领域划分有了更直观的认识，能够根据业务边界拆分用户、商家、交易、聊天与 AI 等模块，并理解网关、配置中心和消息队列在系统中的实际作用。",
    ],
    "成员2收获": [
        "团队在完整交易链路设计上获得了较大提升，能够围绕购物车、订单、支付、库存、发货和售后构建闭环流程，而不是只实现孤立页面或简单 CRUD 接口。",
    ],
    "成员3收获": [
        "团队掌握了前后端协同开发与多角色页面组织方法，能够将用户端、商家端、管理员端和 AI 浮层统一到同一套前端应用中进行组织与联调。",
    ],
    "成员4收获": [
        "团队在分布式一致性和异步解耦方面积累了实践经验，例如通过库存锁定、消息事件、状态枚举和分布式锁来保障交易主链路的稳定性。",
    ],
    "成员5收获": [
        "团队将 AI 服务真正嵌入到订单咨询、售后识别、图搜商品和确认式动作执行等场景中，理解了知识检索、向量库、多 Agent 路由与业务安全边界之间的关系。",
    ],
    "反思": [
        "在完成系统的同时，团队也认识到课程项目从“功能能用”走向“工程成熟”仍有明显提升空间，需要在一致性治理、自动化测试、文档标准化和产品细节方面继续完善。",
    ],
    "成员1反思": [
        "当前系统虽然已经具备完整模块，但部分跨服务链路仍更依赖联调验证而非自动化回归，后续应继续补齐关键接口与主链路测试，提高变更后的稳定性。",
    ],
    "成员2反思": [
        "订单、支付、库存和售后之间的状态管理已经成形，但若面向更高并发或更复杂业务，还需要进一步完善幂等控制、重试策略和异常补偿机制。",
    ],
    "成员3反思": [
        "前端页面入口较完整，但在交互细节、异常提示一致性和移动端适配精细度上仍可继续打磨，使系统更贴近真实产品体验。",
    ],
    "成员4反思": [
        "AI 模块已经具备较强亮点，但模型效果、知识更新流程、提示词治理和风险动作控制仍需长期迭代，不能将一次接入视为最终完成。",
    ],
    "成员5反思": [
        "团队在课程后期才集中整理测试与综合文档，说明过程性文档管理仍不够前置。后续应在开发过程中同步维护设计、测试和总结材料，降低结项阶段的集中整理成本。",
    ],
}

DETAIL_MODULE_2_CONTENT = [
    "AI 客服与售后会话模块负责承接用户在聊天页面中的常规咨询、订单问答、售后意图识别、图片搜商品与确认式动作请求，并与聊天服务、订单服务和商家数据形成联动。模块内部通过 ConversationContext 统一组装文本、图片、订单号与历史对话，再由 AgentRouter 按场景路由到通用问答、售后处理或图像搜索分支。",
    "在实现层面，该模块结合 Spring AI、RAG、Qdrant 向量检索、商品图片索引与 confirmationToken 机制，既保证 AI 输出具备知识依据和上下文理解能力，又避免高风险售后动作被模型直接越权执行，是本项目区别于普通电商课程作品的重要特色模块。",
]


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def insert_body_paragraph_after(paragraph: Paragraph, text: str, template_paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(template_paragraph._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para


def build_heading_anchors(doc: Document) -> dict[str, Paragraph]:
    anchors: dict[str, Paragraph] = {}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        text = paragraph.text.strip()
        if text and style_name.startswith("Heading") and text not in anchors:
            anchors[text] = paragraph
    return anchors


def fill_section(anchors: dict[str, Paragraph], heading_text: str, texts: list[str], template_paragraph: Paragraph) -> None:
    anchor = anchors.get(heading_text)
    if anchor is None:
        raise ValueError(f"Heading not found: {heading_text}")
    current = anchor
    for text in texts:
        current = insert_body_paragraph_after(current, text, template_paragraph)


def update_headers(doc: Document) -> None:
    header_text = f"{PROJECT_NAME}                                       文档编号：{DOC_CODE}"
    for section in doc.sections:
        if section.header.paragraphs:
            set_paragraph_text(section.header.paragraphs[0], header_text)


def fill_cover(doc: Document) -> None:
    set_paragraph_text(doc.paragraphs[6], PROJECT_NAME)
    set_paragraph_text(doc.paragraphs[14], TEAM_NAME)


def fill_team_table(doc: Document) -> None:
    table = doc.tables[0]
    for row_idx, values in enumerate(TEAM_ROWS, start=1):
        for col_idx, value in enumerate(values):
            table.rows[row_idx].cells[col_idx].text = value


def fill_revision_table(doc: Document) -> None:
    table = doc.tables[1]
    for row_idx, values in enumerate(REVISION_ROWS, start=1):
        for col_idx, value in enumerate(values):
            table.rows[row_idx].cells[col_idx].text = value


def fill_function_table(doc: Document) -> None:
    table = doc.tables[2]
    for row_idx, values in enumerate(FUNCTION_TABLE_ROWS, start=1):
        for col_idx, value in enumerate(values):
            table.rows[row_idx].cells[col_idx].text = value
            tc_pr = table.rows[row_idx].cells[col_idx]._tc.get_or_add_tcPr()
            v_merge = tc_pr.find(qn("w:vMerge"))
            if v_merge is not None:
                tc_pr.remove(v_merge)


def fill_use_case_table(doc: Document) -> None:
    table = doc.tables[3]
    for row in table.rows:
        key = row.cells[0].text.strip()
        if key in USE_CASE_TABLE_ROWS:
            row.cells[1].text = USE_CASE_TABLE_ROWS[key]


def update_special_headings(doc: Document) -> None:
    function_headings = []
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if paragraph.text.strip() == "**功能模块" and style_name.startswith("Heading"):
            function_headings.append(paragraph)
    if len(function_headings) != 3:
        raise ValueError("Unexpected number of 功能模块 headings")
    set_paragraph_text(function_headings[0], "用户、商家、交易与智能服务功能模块")
    set_paragraph_text(function_headings[1], "交易履约与购物车模块")
    set_paragraph_text(function_headings[2], "AI客服与售后会话模块")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "表2 ****用例规约":
            set_paragraph_text(paragraph, "表2 用户购物与售后闭环用例规约")
            break


def fill_detail_module_2(doc: Document, template_paragraph: Paragraph) -> None:
    heading = None
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if paragraph.text.strip() == "AI客服与售后会话模块" and style_name.startswith("Heading"):
            heading = paragraph
            break
    if heading is None:
        raise ValueError("Heading not found: AI客服与售后会话模块")
    current = heading
    for text in DETAIL_MODULE_2_CONTENT:
        current = insert_body_paragraph_after(current, text, template_paragraph)


def main() -> None:
    doc = Document(TEMPLATE_PATH)
    body_template = doc.paragraphs[118]
    anchors = build_heading_anchors(doc)

    fill_cover(doc)
    update_headers(doc)
    fill_team_table(doc)
    fill_revision_table(doc)
    fill_function_table(doc)
    fill_use_case_table(doc)
    update_special_headings(doc)

    for heading_text, texts in SECTION_CONTENT.items():
        fill_section(anchors, heading_text, texts, body_template)

    fill_detail_module_2(doc, body_template)

    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
